import os
import base64
import streamlit as st
from langchain.chat_models import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_experimental.text_splitter import SemanticChunker
from langchain_core.runnables import RunnablePassthrough
from langchain.schema import StrOutputParser
from langchain_community.embeddings import HuggingFaceBgeEmbeddings,HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from docx import Document
# Set up
import time
from langchain.prompts import PromptTemplate

import logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.ERROR)
if not logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)

os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_PROJECT"]="test1"
os.environ["LANGCHAIN_ENDPOINT"]="https://api.smith.langchain.com"
os.environ["LANGCHAIN_API_KEY"]=""
os.environ["LANGCHAIN_API_KEY"]="ls__fd390e70390e4be2af0baf6b10599d91"

def main_bg(main_bg):
    main_bg_ext = "png"
    st.markdown(
        f"""
        <div class='header-area'>
                <img src="data:image/{main_bg_ext};base64,{base64.b64encode(open(main_bg, "rb").read()).decode()}"/>
                </div>
         <style>
         .header-area {{
            height:60px !important;
            background-color:black !important;
            width:100vw;
            position:absolute;
            margin:0 auto;
            margin-top:-30px;
            position: fixed;
            margin-top: -80px;
            left: 0;
            padding-left:50px;
            z-index: 999999;
            img{{
              height:60px;
            }}
         }}
         </style>
         """,
        unsafe_allow_html=True
    )
main_bg('./img/lion.png')
def clear_api_base():
  env = dict(os.environ)
  if "OPENAI_API_BASE" in env:
    os.environ.pop("OPENAI_API_BASE")
    os.environ["OPENAI_API_KEY"] = openaikey
  else:
    return

template = """
##Capacity and Role##
你是一个资深的律师，十分擅长民事领域的咨询问答

##Context##
请充分理解我给你的资深律师的办案心得和示例，深入学习示例的解答逻辑和语气风格。
严格依据民法典实施后仍然有效的法律和司法解释，用准确、清晰、简明扼要且符合法律逻辑的语言实质性解答客户提出的法律咨询。

请根据下面的Question, Examples, 及History, 提供实质性答案
其中每个元素的说明如下:
Question: 用户提出的问题,需要基于此问题进行解答
Example: 根据用于问题查找到的相关资料，但是注意这些资料未必与用户的问题完全匹配，请根据实际情况参考这些资料
History: 与该用户的历史对话，可以参考

##Question##
{question}

##Examples##
{examples}

##History##
{chat_history}

##Output Indicator##
    1.作为资深律师，请以法律逻辑的语言为客户的法律咨询提供实质性答案。
    2.每个回答都需要根据"提供的示例"来回答，如果不匹配，则建议咨询律师。不要提供专业律师的建议。
    3.请尽量简明扼要，不要提供过多的信息，以免造成用户的困惑。
    4.请将每次回答的字数控制在100字以内。

"""

embeddings1 = HuggingFaceEmbeddings(model_name=os.getenv('EMB_MODEL_PATH', '/models/bge-small-zh-v1.5'))
        


# 初始化文本分割器
text_splitter1 = RecursiveCharacterTextSplitter(
    chunk_size=2000, 
    chunk_overlap=20,
    length_function = len
  )
split_doc_threshold = '50'
text_splitter2 = SemanticChunker(embeddings1,
    breakpoint_threshold_type="percentile", 
    breakpoint_threshold_amount=float(split_doc_threshold)) #语义分割,0.5

def load_local_file():
    text = ""
    docx = Document('./docs/law.docx')
    for paragraph in docx.paragraphs:
          text += paragraph.text
    print('docx文档处理完毕')
    return text


def load_file():
    documents = []
    text = ''
    file_dir = './docs/law.docx'
    docx = Document(file_dir)
    for paragraph in docx.paragraphs:
          text = paragraph.text
          documents.append({"text": text, "metadata": {"source": file_dir}})
    texts = [doc["text"] for doc in documents]
    # print(documents,'docs')
    # print(texts,'文本内容')
    return texts
# load_file()
    # return text
 
def load_retriever(text,emb):
    model_name = emb
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': True} # set True to compute cosine similarity
    embeddings = HuggingFaceBgeEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs,
        query_instruction="为这个句子生成表示以用于检索相关文章："
        )
    kb_dir = f"tmp_vdb"
    os.makedirs(kb_dir, exist_ok=True)
    logger.debug("Splitting text into chunks. beging...")
    print('构造知识库')


    # 使用文本切割
    # chunks = text_splitter1.split_text(text)
    # knowledge_base = FAISS.from_texts(chunks, embeddings)

    # 使用语义切割
    texts = load_file()
    docs = text_splitter2.create_documents(texts)
    knowledge_base = FAISS.from_documents(docs, embeddings1)

    knowledge_base.save_local(kb_dir)
    print('构造完毕')
    return knowledge_base

def clear_retriever():

  if os.path.exists('./tmp_vdb/index.faiss'): # 存在缓存，直接读取
    os.remove('./tmp_vdb/index.faiss')
    os.remove('./tmp_vdb/index.pkl')
    st.success('成功清除缓存', icon="✅")
  else:
    st.info('暂无缓存，请先上传文件', icon="ℹ️")
    
doc_data = ''
def make_qa_chain(text, chain_type, k, history, llm, emb,tem,q):
    print('开始构建c')
    emb_t1 = time.time()
    if os.path.exists('./tmp_vdb/index.faiss'): # 存在缓存，直接读取
      model_name = emb
      model_kwargs = {'device': 'cpu'}
      encode_kwargs = {'normalize_embeddings': True} # set True to compute cosine similarity

      embeddings = HuggingFaceBgeEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs,
        query_instruction="为这个句子生成表示以用于检索相关文章："
        )
      # 使用文本分割向量
      # knowledge_base = FAISS.load_local("./tmp_vdb",embeddings,allow_dangerous_deserialization=True)
      # 使用语义分割
      knowledge_base = FAISS.load_local("./tmp_vdb",embeddings,allow_dangerous_deserialization=True)
    else:
      knowledge_base = load_retriever(text,emb) # 不存在缓存，调用接口
    
    retriever = knowledge_base.as_retriever(search_type="similarity", search_kwargs={"k": k})
    # retriever = knowledge_base.similarity_search_with_score(q)
    print(retriever,'retriever')

    emb_t2 = time.time()
    print('检索召回内容',emb_t2 - emb_t1)

    chat_t1 = time.time()
    ## 通过 LCEL 组装chat

    def format_docs(docs):
        global doc_data
        doc_data = [d.page_content for d in docs]
        doc_data = "".join(doc_data)
        doc_data = doc_data.replace('\u3000','<br/>')
        return [d.page_content for d in docs]


    def format_history(history):
        # return "\n\n".join(history)
        return history

    llm = ChatOpenAI(
        streaming=True,
        verbose=True,
        openai_api_key=os.getenv("OPENAI_API_KEY", "none"),
        openai_api_base = os.getenv("OPENAI_BASE_URL", "http://172.16.30.218:8000/v1"),
        model_name=os.getenv("OPENAI_MODEL", "Qwen/Qwen2-7B-Instruct")
    )

    QA_CHAIN_PROMPT = PromptTemplate(input_variables=["examples", "question", "chat_history"],template=tem)
    qa_chain = (
        {
            "chat_history": (lambda _: format_history(history)), 
            'examples': retriever | format_docs, #使用retriever召回topk的文档
            "question": RunnablePassthrough()
        }
        | QA_CHAIN_PROMPT # prompt模板
        | llm # 所选模型
        | StrOutputParser() # 输入输出的格式化
    )

    
    chat_t2 = time.time()
    print('调用chain耗时',chat_t2 - chat_t1)
    return qa_chain


chat_history=[]

def main():

    # llm_name = st.sidebar.selectbox(
    #   "请选择模型",
    # ("Qwen/Qwen2-7B-Instruct",)
    # )
    # emb_name = st.sidebar.selectbox(
    #   "请选择模型",
    # ("BAAI/bge-small-zh-v1.5","BAAI/bge-base-zh-v1.5")
    # )
    # prompt_template = st.sidebar.text_area(
    # label = "Template",
    # value = template,
    # height = 200,
    # )
    # tar_file = st.sidebar.file_uploader("选择一个文件",accept_multiple_files=True,type=["pdf","docx","doc"])
    # llm_name = 'Qwen/Qwen2-7B-Instruct'
    # emb_name = 'BAAI/bge-small-zh-v1.5'
    # if st.sidebar.button('生成知识库'):
    #     text = load_local_file()
    #     load_retriever(text,emb=emb_name)
    #     print('生成知识库')
    # st.sidebar.button('清空缓存', on_click=clear_retriever)
    chat_history=[]
    text =''
    st.header('某律师事务所法务助手', divider='rainbow')
    #llm_name = 'Qwen/Qwen2-7B-Instruct'
    llm_name=os.getenv("OPENAI_MODEL", "Qwen/Qwen2-7B-Instruct")
    emb_name = os.getenv('EMB_MODEL_PATH', '/models/bge-small-zh-v1.5')
    prompt_template = template
    if "messages" not in st.session_state:
        st.session_state.messages = []
        st.session_state.messages.append({"role": "assistant", "content": '您好，我是某律师事务所法务助手，请问有什么能帮助您的？'})
    for message in st.session_state.messages:
        with st.chat_message(message["role"]): 
            st.markdown(message["content"])
    if prompt := st.chat_input("请输入您的问题"): 
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                full_response = ""
                chat_history = st.session_state.messages
                print(chat_history,'聊天记录')
                qa_chain = make_qa_chain(text, "stuff", 4, history=chat_history, llm=llm_name, emb=emb_name,tem = prompt_template,q=prompt)
                qa_t1 = time.time()
                result = qa_chain.stream(prompt)
                qa_t2 = time.time()
                print('输入问题到作答耗时',qa_t2 - qa_t1)             
                assistant_response =result
                for chunk in assistant_response:
                    full_response += chunk
                    message_placeholder.markdown(full_response + "▌")
                message_placeholder.markdown('''{}'''.format(full_response),unsafe_allow_html=True)
                with st.expander("参考资料"):
                    st.markdown(doc_data,unsafe_allow_html=True)
                st.session_state.messages.append({"role": "assistant", "content": full_response})
if __name__ == '__main__':
    # set_api_key()
    main()
